#! /user/bin/python3

#  Copyright (c) 2024. All rights reserved.
#  This source code is licensed under the CC BY-NC-SA
#  (Creative Commons Attribution-NonCommercial-NoDerivatives) License, By Xiao Songtao.
#  This software is protected by copyright law. Reproduction, distribution, or use for commercial
#  purposes is prohibited without the author's permission. If you have any questions or require
#  permission, please contact the author: 2207150234@st.sziit.edu.cn

# -------------------------<edocsitahw>----------------------------
# 传建时间: 2024/6/28 上午8:20
# 当前项目名: Python
# 编码模式: utf-8
# 注释: 
# -------------------------<edocsitahw>----------------------------
from abc import ABC, abstractmethod
from os import PathLike, path
from PyQt6.QtWidgets import QWidget, QLabel, QVBoxLayout, QPushButton
from PyQt6.QtGui import QFont, QColor, QPalette
from PyQt6.QtCore import QPropertyAnimation, pyqtSlot
from docx import Document
from docx.text.paragraph import Paragraph
from re import findall
from pandas import DataFrame, Series
from warnings import warn
from typing import final
from random import choice
from functools import partial
from customWidget import PushButton
from time import sleep


class QuestionPool:
    def __init__(self, data: DataFrame, *, Type: type["base"] = None):
        self._data = data
        self._Type = Type
        self._qPool = []
        self._idx = 1
        self._len = len(self._data)
        self.fillPool()

    @property
    def data(self) -> DataFrame:
        return self._data

    def fillPool(self):
        for i, d in self._data.iterrows():
            self._qPool.append(d.to_dict())

    def __len__(self):
        return self._len

    def __iter__(self):
        return self

    def __next__(self) -> dict:
        if self._qPool:
            questionDict = self._qPool.pop(choice(range(len(self._qPool))))

            if self._Type:
                infoDict = self._Type().infoDict

                for k, v in questionDict.items():
                    if not isinstance(v, infoDict[k]):
                        try:
                            if infoDict[k] == list:
                                tf = lambda x: f'"{x}"'
                                questionDict[k] = f"[{','.join([tf(_) for _ in v.replace('[', '').replace(']', '').replace(',', '')])}]"

                            questionDict[k] = eval(questionDict[k])
                        except NameError as e:
                            raise NameError(f"{self._Type.__name__}: {infoDict} {v}({type(v).__name__}) -> {infoDict[k]} {e}!")

            if "id" in questionDict:
                questionDict["id"] = self._idx
                self._idx += 1

            self._len -= 1

            return questionDict
        else:
            raise StopIteration


class Parser:
    def __init__(self, docxPath: str | PathLike[str], csvDirPath: str | PathLike[str]):
        self._docxPath = docxPath
        self._csvDirPath = csvDirPath
        self._document = Document(docxPath)

        if not path.exists(csvDirPath):
            raise FileNotFoundError(
                f"{csvDirPath}不存在!")

        if not path.isdir(csvDirPath):
            raise NotADirectoryError(
                f"{csvDirPath}应该是一个文件夹!")

    @property
    def document(self):
        return self._document

    def parsing(self, *args: type["base"]):
        for pType in args:
            if not issubclass(pType, base):
                raise TypeError(
                    f"Argument {pType} is not a subclass of {base}")

        try:
            parser = None
            for i, para in enumerate(self.document.paragraphs):

                for pType in args:
                    if pType.startSignal(para):

                        parser = pType()

                    if pType.endSignal(para, self.document.paragraphs[i + 1]):
                        parser.save(path.join(self._csvDirPath, (parser.name[0] if isinstance(parser.name, list) else parser.name) + ".csv"))

                if parser is not None:
                    parser.parser(para)

        except IndexError as e:
            warn(
                f"IndexError: {e}")


class base(ABC):
    name = "base"

    def __init__(self):
        self._data = DataFrame(columns=list(self.infoDict.keys()))
        self._tmpData = {}

    @property
    def data(self) -> DataFrame:
        return self._data

    @property
    def cnName(self) -> str:
        return self.name

    @property
    @abstractmethod
    def infoDict(self) -> dict[str, type]:
        pass

    @final
    def register(self, **kwargs):
        for k, v in kwargs.items():
            if not isinstance(v, self.infoDict[k]):
                raise TypeError(
                    f"Value of {k} should be {self.infoDict[k].__name__}!")

            if k not in self.infoDict:
                return KeyError(
                    f"key {k} is not in {self.infoDict}")

            self._tmpData[k] = v

        if all(k in self._tmpData for k in self.infoDict):

            self._data = self._data._append(Series(self._tmpData), ignore_index=True)

            self._tmpData = {}

    @abstractmethod
    def parser(self, paragraph: Paragraph):
        pass

    @abstractmethod
    def showUI(self, parent: QWidget, data: DataFrame):
        pass

    @final
    def save(self, csvPath: str | PathLike[str]):
        self._data.to_csv(csvPath, index=False, encoding="gbk")

    @classmethod
    @abstractmethod
    def startSignal(cls, paragraph: Paragraph) -> bool:
        pass

    @classmethod
    @abstractmethod
    def endSignal(cls, paragraph: Paragraph, lastPara: Paragraph) -> bool:
        pass


class Choice(base):
    name = "单选题"

    def __init__(self):
        super().__init__()
        self._question = ""
        self._idx = 0
        self._option = {}

    @property
    def infoDict(self) -> dict[str, type]:
        return {

            "id": int,
            "question": str,
            "option": dict,
            "answer": str
        }

    def parser(self, paragraph: Paragraph):
        text = paragraph.text.replace("\t", "")

        try:
            if text[0].isdigit():
                res = findall(r"(\d{1,2}).(?!=\s)(.*)$", text)
                idx, self._question = int(res[0][0]), res[0][1]

                if idx != self._idx + 1:
                    warn(
                        f"The id of question should be {self._idx+1} but {idx} found!")

                self._idx = idx

                self.register(id=idx, question=self._question)

            elif text[0] in ["A", "B", "C", "D"]:
                res = findall(r"([ABCD]).(?!=\s)(.*)$", text)
                option, optionText = res[0][0], res[0][1]

                self._option[option] = optionText

            if len(self._option) == 4:
                self.register(option=self._option, answer="")

                self._option, self._question = {}, ""

        except Exception as e:
            warn(
                f"Error occurred while parsing question {self._idx+1}: {e}")

    def btnClick(self, pool: QuestionPool, questionLabel: QLabel, parent: QWidget, topLayout: QVBoxLayout, answerLabel: QLabel, *, option: str, btn: QLabel, answer: str):
        # animation = QPropertyAnimation(btn, b"color")
        # animation.setDuration(2500)
        # animation.setStartValue(btn.originColor)
        # animation.setEndValue(QColor(255, 0, 0))
        # animation.finished.connect(lambda: self.animFinished(pool, questionLabel, parent, topLayout))
        # animation.start()

        btn.setText(f"上一题答案: {answer}, 你的答案是: {option}!")

        if len(pool) > 0:
            self.nextQuestion(pool, questionLabel, parent, topLayout, answerLabel)

        return
    # @pyqtSlot()
    # def animFinished(self, pool: QuestionPool, qLabel: QLabel, parent: QWidget, topLayout: QVBoxLayout):
    #     self.nextQuestion(pool, qLabel, parent, topLayout)

    def showUI(self, parent: QWidget, data: DataFrame):
        qPool = QuestionPool(data)

        topLayout = QVBoxLayout(parent)

        questionLabel = QLabel(parent)

        topLayout.addWidget(questionLabel)

        font = QFont()
        font.setPointSize(20)

        questionLabel.setFont(font)

        questionLabel.resize(parent.width(), 40)

        answerLabel = QLabel(parent)

        topLayout.addWidget(answerLabel)

        if len(qPool) > 0:
            self.nextQuestion(qPool, questionLabel, parent, topLayout, answerLabel)

    def nextQuestion(self, pool: QuestionPool, qLabel: QLabel, parent: QWidget, topLayout: QVBoxLayout, answerLabel: QLabel):
        if len(pool) > 0:

            qDict = next(pool)

            btnDict = {}

            qLabel.setText(f"{qDict['id']}. {qDict['question']}")

            self.clearButtons(topLayout)

            for i, option in eval(qDict["option"]).items():
                if i not in btnDict:
                    # btnDict[i] = PushButton(parent)
                    btnDict[i] = QPushButton(parent)
                    btnDict[i].setObjectName(i)
                    btnDict[i].setStyleSheet("QPushButton { text-align: left;}")# border: none; background-color: transparent; }")

                btnDict[i].setText(f"{i}. {option}")
                clickFunc = partial(self.btnClick, pool, qLabel, parent, topLayout, answerLabel, option=i, btn=answerLabel, answer=qDict["answer"])
                btnDict[i].clicked.connect(clickFunc)
                topLayout.addWidget(btnDict[i])

    @staticmethod
    def clearButtons(topLayout: QVBoxLayout):
        for i in reversed(range(topLayout.count())):
            widget = topLayout.takeAt(i).widget()
            if isinstance(widget, QPushButton):
                widget.setParent(None)

    @classmethod
    def startSignal(cls, paragraph: Paragraph):
        return True if cls.name in paragraph.text else False

    @classmethod
    def endSignal(cls, paragraph: Paragraph, lastPara: Paragraph):
        return True if "多项选择题" in lastPara.text else False


class ChoiceMulti(base):
    name = ["多项选择题", "多选题"]

    def __init__(self):
        super().__init__()
        self._question = ""
        self._idx = 0
        self._option = {}

    @property
    def cnName(self) -> list[str]:
        return ["多项选择题", "多选题"]

    @property
    def infoDict(self) -> dict[str, type]:
        return {
            "id": int,
            "question": str,
            "option": dict,
            "answer": list
        }

    def parser(self, paragraph: Paragraph):
        text = paragraph.text.replace("\t", "")

        try:
            if text[0].isdigit():
                res = findall(r"(\d{1,2}).(?!=\s)(.*)$", text)
                idx, self._question = int(res[0][0]), res[0][1]

                if idx != self._idx + 1:
                    warn(
                        f"The id of question should be {self._idx + 1} but {idx} found!")

                self._idx = idx

                self.register(id=idx, question=self._question)

            elif text[0] in ["A", "B", "C", "D", "E"]:
                res = findall(r"([ABCDE]).(?!=\s)(.*)$", text)
                option, optionText = res[0][0], res[0][1]

                self._option[option] = optionText

            if len(self._option) == 5:
                self.register(option=self._option, answer=[])

                self._option, self._question = {}, ""

        except Exception as e:
            warn(
                f"Error occurred while parsing question {self._idx+1}: {e}")

    def showUI(self, parent: QWidget, data: DataFrame):
        pass

    @classmethod
    def startSignal(cls, paragraph: Paragraph):
        return True if any(n in paragraph.text for n in cls.name) else False

    @classmethod
    def endSignal(cls, paragraph: Paragraph, lastPara: Paragraph):
        return True if "判断题" in lastPara.text else False


class Judge(base):
    name = "判断题"

    def __init__(self):
        super().__init__()
        self._question = ""
        self._idx = 0

    @property
    def infoDict(self) -> dict[str, type]:
        return {
            "id": int,
            "question": str,
            "option": dict,
            "answer": str,
        }

    def parser(self, paragraph: Paragraph):
        text = paragraph.text

        try:
            if text[0].isdigit():
                res = findall(r"(\d{1,2}).(?!=\s)(.*)(?=[(（])", text)

                idx, self._question = int(res[0][0]), res[0][1]

                if idx != self._idx + 1:
                    warn(
                        f"The id of question should be {self._idx + 1} but {idx} found!")

                self._idx = idx

                self.register(id=idx, question=self._question, answer="")

        except Exception as e:
            warn(
                f"Error occurred while parsing question {self._idx+1}: {e}")

    def showUI(self, parent: QWidget, data: DataFrame):
        pass

    @classmethod
    def startSignal(cls, paragraph: Paragraph) -> bool:
        return True if cls.name in paragraph.text else False

    @classmethod
    def endSignal(cls, paragraph: Paragraph, lastPara: Paragraph) -> bool:
        return True if "答案" in lastPara.text else False


if __name__ == '__main__':
    # par = Parser(r"E:\codeSpace\codeSet\Python\privateProject\exam\static\doc\test.docx", r"E:\codeSpace\codeSet\Python\privateProject\exam\static\csv")
    # par.parsing(Choice, ChoiceMultiple, Judge)
    # from pandas import read_csv
    # df = read_csv(r"E:\codeSpace\codeSet\Python\privateProject\exam\static\csv\单选题.csv", encoding='gbk')
    # pool = QuestionPool(df)
    # pool.fillPool()
    par = Parser(r"C:\Users\Lenovo\Desktop\考试复习\云计算题库46题.docx", r"E:\codeSpace\codeSet\Python\privateProject\exam\static\csv")
    par.parsing(Judge)
